import os
import re
from typing import Any, Callable
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils import copiar_a_temp, fecha_texto, limpiar_nombre

PDF_DISPONIBLE: bool = False
try:
    from docx2pdf import convert as _pdf_conv
    PDF_DISPONIBLE = True
except ImportError:
    try:
        import subprocess as _sp
        _sp.run(["libreoffice", "--version"], capture_output=True, check=True)
        PDF_DISPONIBLE = True
    except (FileNotFoundError, _sp.CalledProcessError):
        PDF_DISPONIBLE = False

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _convertir_a_pdf(ruta_docx: str) -> None:
    ruta_pdf = os.path.splitext(ruta_docx)[0] + ".pdf"
    try:
        from docx2pdf import convert as _conv
        _conv(ruta_docx, ruta_pdf)
    except ImportError:
        import subprocess as _sp
        _sp.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
             os.path.dirname(ruta_pdf), ruta_docx],
            capture_output=True, check=True, timeout=60,
        )


def _reformatear_si_fecha(valor: str, fmt_destino: str) -> str | None:
    """Reformatea `valor` como fecha si se puede parsear. Retorna None si no."""
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S",
                "%d/%m/%y", "%d-%m-%y", "%Y/%m/%d", "%d.%m.%Y", "%d.%m.%y"):
        try:
            dt = datetime.strptime(valor.strip(), fmt)
            return fecha_texto(dt) if fmt_destino == "texto" else dt.strftime(fmt_destino)
        except ValueError:
            continue
    return None


def _extraer_formato_run(r) -> dict[str, Any]:
    """Extrae bold/italic/underline/size/name de un elemento <w:r> XML."""
    rPr = r.find(W_NS + "rPr")
    fmt: dict[str, Any] = {
        "bold":      False,
        "italic":    False,
        "underline": False,
        "size":      None,
        "name":      None,
    }
    if rPr is not None:
        fmt["bold"]      = rPr.find(W_NS + "b") is not None
        fmt["italic"]    = rPr.find(W_NS + "i") is not None
        fmt["underline"] = rPr.find(W_NS + "u") is not None
        sz = rPr.find(W_NS + "sz")
        if sz is not None:
            fmt["size"] = sz.get(qn("w:val"))
        rFonts = rPr.find(W_NS + "rFonts")
        if rFonts is not None:
            fmt["name"] = rFonts.get(qn("w:ascii"))
    return fmt


def _crear_run_xml(texto: str, fmt: dict[str, Any] | None):
    """Crea un elemento <w:r> con su <w:t> y formato."""
    r = OxmlElement("w:r")
    if fmt and (fmt["bold"] or fmt["italic"] or fmt["underline"] or fmt["size"] or fmt["name"]):
        rPr = OxmlElement("w:rPr")
        if fmt.get("bold"):
            rPr.append(OxmlElement("w:b"))
        if fmt.get("italic"):
            rPr.append(OxmlElement("w:i"))
        if fmt.get("underline"):
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        if fmt.get("size"):
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(fmt["size"]))
            rPr.append(sz)
        if fmt.get("name"):
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), fmt["name"])
            rFonts.set(qn("w:hAnsi"), fmt["name"])
            rPr.append(rFonts)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = texto
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _mismo_formato(a: dict[str, Any] | None, b: dict[str, Any] | None) -> bool:
    if a is None or b is None:
        return a is b
    return (
        a["bold"] == b["bold"]
        and a["italic"] == b["italic"]
        and a["underline"] == b["underline"]
        and a["size"] == b["size"]
        and a["name"] == b["name"]
    )


def _dividir_texto_por_formato(
    texto: str, inicio: int, fin: int, mapa_formato: list[dict[str, Any]],
) -> list[tuple[str, dict[str, Any] | None]]:
    """Divide ``texto[inicio:fin]`` en sub-segmentos de formato uniforme."""
    if inicio >= fin:
        return []
    segmentos: list[tuple[str, dict[str, Any] | None]] = []
    i = inicio
    while i < fin:
        fmt = mapa_formato[i] if i < len(mapa_formato) else None
        j = i + 1
        while j < fin and j < len(mapa_formato) and _mismo_formato(mapa_formato[j], fmt):
            j += 1
        segmentos.append((texto[i:j], fmt))
        i = j
    return segmentos


def _generar_segmentos(texto_completo: str, runs: list) -> list[tuple[str, dict[str, Any] | None]]:
    """Construye segmentos (texto_sin_reemplazar, formato) a partir del texto completo y sus runs.

    El texto de los segmentos placeholder se deja sin reemplazar (ej. ``{{nombre}}``);
    el llamador debe aplicar el reemplazo después.
    """
    mapa_formato: list[dict[str, Any]] = []
    for r in runs:
        fmt = _extraer_formato_run(r)
        text_len = sum(len(t.text or "") for t in r.iter(W_NS + "t"))
        mapa_formato.extend(fmt for _ in range(text_len))

    reemplazos_formato: dict[str, Any] = {}
    for m in re.finditer(r"\{\{([^}]+)\}\}", texto_completo):
        campo = m.group(1).strip()
        pos = m.start()
        reemplazos_formato[campo.lower()] = mapa_formato[pos] if pos < len(mapa_formato) else None

    segmentos: list[tuple[str, Any]] = []
    cursor = 0
    for m in re.finditer(r"\{\{([^}]+)\}\}", texto_completo):
        if m.start() > cursor:
            segmentos.extend(
                _dividir_texto_por_formato(texto_completo, cursor, m.start(), mapa_formato)
            )
        segmentos.append((
            m.group(0),
            reemplazos_formato.get(m.group(1).strip().lower()),
        ))
        cursor = m.end()
    if cursor < len(texto_completo):
        segmentos.extend(
            _dividir_texto_por_formato(texto_completo, cursor, len(texto_completo), mapa_formato)
        )
    return segmentos



def generar_documentos(
    word_path: str,
    trabajadores: list[dict[str, Any]],
    columnas: list[str],
    formato_fecha: str,
    col_identificador: str,
    campos_extra_vars: dict[str, dict[str, Any]],
    formato_nombre: str,
    usar_guion_bajo: bool,
    carpeta_destino: str,
    generar_pdf: bool = False,
    on_progreso: Callable[[int, int], None] | None = None,
) -> tuple[int, list[str]]:
    errores: list[str] = []
    generados = 0
    total = len(trabajadores)

    def reemplazar_texto(texto: str) -> str:
        def sub(m: re.Match) -> str:
            campo = m.group(1).strip()
            for k, v in reemplazos.items():
                if k.lower() == campo.lower():
                    return v
            return m.group(0)
        return re.sub(r"\{\{([^}]+)\}\}", sub, texto)

    def procesar_parrafo(parrafo: Any) -> None:
        texto_completo = "".join(r.text for r in parrafo.runs)
        nuevo_texto = reemplazar_texto(texto_completo)
        if texto_completo == nuevo_texto:
            return

        mapa_formato: list[dict[str, Any]] = []
        for r in parrafo.runs:
            for _ in r.text:
                mapa_formato.append({
                    "bold":      r.bold,
                    "italic":    r.italic,
                    "underline": r.underline,
                    "size":      r.font.size,
                    "name":      r.font.name,
                })

        reemplazos_formato: dict[str, Any] = {}
        for m in re.finditer(r"\{\{([^}]+)\}\}", texto_completo):
            campo = m.group(1).strip()
            pos = m.start()
            fmt = mapa_formato[pos] if pos < len(mapa_formato) else None
            reemplazos_formato[campo.lower()] = fmt

        segmentos: list[tuple[str, Any]] = []
        cursor = 0
        for m in re.finditer(r"\{\{([^}]+)\}\}", texto_completo):
            campo = m.group(1).strip()
            if m.start() > cursor:
                segmentos.extend(
                    _dividir_texto_por_formato(texto_completo, cursor, m.start(), mapa_formato)
                )
            valor = reemplazar_texto(m.group(0))
            fmt = reemplazos_formato.get(campo.lower())
            segmentos.append((valor, fmt))
            cursor = m.end()
        if cursor < len(texto_completo):
            segmentos.extend(
                _dividir_texto_por_formato(texto_completo, cursor, len(texto_completo), mapa_formato)
            )

        run_base = parrafo.runs[0] if parrafo.runs else None
        for r in parrafo.runs:
            r.text = ""

        for texto_seg, fmt in segmentos:
            if not texto_seg:
                continue
            run = parrafo.add_run(texto_seg)
            if fmt:
                run.bold      = fmt["bold"]
                run.italic    = fmt["italic"]
                run.underline = fmt["underline"]
                run.font.size = fmt["size"]
                if fmt["name"]:
                    run.font.name = fmt["name"]
            elif run_base:
                run.bold      = run_base.bold
                run.italic    = run_base.italic
                run.underline = run_base.underline
                run.font.size = run_base.font.size

    def _reemplazar_runs_xml(runs: list, texto_completo: str) -> None:
        """Reemplaza {{placeholders}} en runs XML preservando el formato run a run."""
        nuevo_texto = reemplazar_texto(texto_completo)
        if nuevo_texto == texto_completo:
            return
        segmentos = _generar_segmentos(texto_completo, runs)
        parent = runs[0].getparent() if runs else None
        if parent is None:
            return
        for r in runs:
            parent.remove(r)
        for texto_seg, fmt in segmentos:
            if not texto_seg:
                continue
            r = _crear_run_xml(reemplazar_texto(texto_seg), fmt)
            parent.append(r)

    for idx, persona in enumerate(trabajadores):
        ruta_lectura: str | None = None
        try:
            ruta_lectura = copiar_a_temp(word_path)
            doc = Document(ruta_lectura)

            reemplazos: dict[str, str] = {k.strip(): str(v) if v is not None else "" for k, v in persona.items()}
            for k, v in persona.items():
                if isinstance(v, datetime):
                    if formato_fecha == "texto":
                        reemplazos[k.strip()] = fecha_texto(v)
                    else:
                        reemplazos[k.strip()] = v.strftime(formato_fecha)
                elif isinstance(v, str):
                    reformateada = _reformatear_si_fecha(v, formato_fecha)
                    if reformateada is not None:
                        reemplazos[k.strip()] = reformateada

            nombre_persona = str(persona.get(col_identificador, str(list(persona.values())[0])))
            extras = campos_extra_vars.get(nombre_persona, {})
            for k, var in extras.items():
                reemplazos[k.strip()] = var.get()

            for p in doc.paragraphs:
                procesar_parrafo(p)
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for p in celda.paragraphs:
                            procesar_parrafo(p)
            for seccion in doc.sections:
                for p in seccion.header.paragraphs:
                    procesar_parrafo(p)
                for p in seccion.footer.paragraphs:
                    procesar_parrafo(p)
            for txbx in doc.element.body.iter(W_NS + "txbxContent"):
                for txbx_p in txbx.iter(W_NS + "p"):
                    runs = list(txbx_p.iter(W_NS + "r"))
                    texto_completo = "".join(
                        "".join(t.text or "" for t in r.iter(W_NS + "t"))
                        for r in runs
                    )
                    if not re.search(r"\{\{[^}]+\}\}", texto_completo):
                        continue
                    _reemplazar_runs_xml(runs, texto_completo)

            nombre_archivo = formato_nombre
            for k, v in reemplazos.items():
                nombre_archivo = nombre_archivo.replace(f"{{{k}}}", str(v))
            nombre_archivo = nombre_archivo.replace("{fecha}", datetime.today().strftime("%Y%m%d"))
            nombre_archivo = limpiar_nombre(nombre_archivo, usar_guion_bajo)
            nombre_archivo = nombre_archivo + ".docx"

            ruta_final = os.path.join(carpeta_destino, nombre_archivo)
            doc.save(ruta_final)
            if generar_pdf:
                try:
                    _convertir_a_pdf(ruta_final)
                except Exception as e:
                    errores.append(f"{nombre_persona}: PDF - {e}")
            generados += 1

        except Exception as e:
            nombre = str(persona.get(col_identificador, str(list(persona.values())[0])))
            errores.append(f"{nombre}: {e}")
        finally:
            if ruta_lectura and os.path.exists(ruta_lectura):
                os.unlink(ruta_lectura)

        if on_progreso:
            on_progreso(idx + 1, total)

    return generados, errores
