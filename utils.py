import os
import re
import tempfile
import shutil
from datetime import datetime
from styles import MESES


def fecha_texto(dt: datetime) -> str:
    return f"{dt.day} de {MESES[dt.month - 1]} del {dt.year}"


def reformatear_fecha(valor: str, fmt_destino: str) -> str:
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y",
                "%d-%m-%y", "%Y/%m/%d", "%d.%m.%Y", "%d.%m.%y"):
        try:
            dt = datetime.strptime(valor.strip(), fmt)
            return fecha_texto(dt) if fmt_destino == "texto" else dt.strftime(fmt_destino)
        except ValueError:
            continue
    return fecha_texto(datetime.today()) if fmt_destino == "texto" else datetime.today().strftime(fmt_destino)


def copiar_a_temp(ruta_original: str) -> str:
    ext = os.path.splitext(ruta_original)[1]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    tmp.close()
    shutil.copy2(ruta_original, tmp.name)
    return tmp.name


def detectar_campos_word(ruta_word: str) -> list[str]:
    if not ruta_word:
        return []
    try:
        from docx import Document
        doc = Document(ruta_word)
        texto = " ".join(p.text for p in doc.paragraphs)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    texto += " " + celda.text
        patron = re.compile(r"\{\{([^}]+)\}\}")
        return list(dict.fromkeys(m.group(1).strip() for m in patron.finditer(texto)))
    except Exception:
        return []


def limpiar_nombre(nombre: str, usar_guion_bajo: bool) -> str:
    sep = "_" if usar_guion_bajo else " "
    nombre = re.sub(r'[\\/*?:"<>|]', sep, nombre)
    nombre = re.sub(r"[ _]+", sep, nombre)
    nombre = nombre.strip(sep)
    if not usar_guion_bajo:
        nombre = nombre.replace("_", " ")
    return nombre
